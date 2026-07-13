"""
Excel/CSV 工具集。
所有函数接收 file（UploadedFile 或文件路径），返回字符串供 LLM 阅读。
"""

import io
import json
import os
import re
import subprocess
import tempfile
import threading
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import matplotlib
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches

matplotlib.use("Agg")
matplotlib.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"]
matplotlib.rcParams["axes.unicode_minus"] = False
import matplotlib.pyplot as plt

# 文件读取锁：防止多个线程同时读取同一个 UploadedFile 对象
_FILE_READ_LOCK = threading.Lock()

# 【关键】模块级文件缓存 —— 跨线程共享，替代 session_state
# ThreadPoolExecutor 的 worker 线程访问不到 st.session_state（它用了 threading.local）
# 但模块级变量对所有线程都可见
_FILE_CACHE: Dict[str, object] = {}  # {"bytes": ..., "name": ...}


def _get_cached_file_data() -> Tuple[Optional[bytes], Optional[str]]:
    """获取预读缓存的文件数据。先查模块缓存，再查 session_state（兼容）。"""
    # 模块级缓存（跨线程可见）
    cached = _FILE_CACHE.get("bytes")
    if cached:
        return cached, _FILE_CACHE.get("name")
    # 回退：session_state 缓存（仅主线程可见，但作为安全网）
    cached_bytes = st.session_state.get("_cached_file_bytes")
    cached_name = st.session_state.get("_cached_file_name")
    if cached_bytes is not None and cached_name is not None:
        return cached_bytes, cached_name
    return None, None


def _set_cached_file_data(data: bytes, name: str):
    """同时更新模块缓存和 session_state 缓存。"""
    _FILE_CACHE["bytes"] = data
    _FILE_CACHE["name"] = name
    try:
        st.session_state["_cached_file_bytes"] = data
        st.session_state["_cached_file_name"] = name
    except Exception:
        pass  # 非 Streamlit 环境下 st.session_state 不可用


def _clear_cached_file_data():
    """清除所有缓存。"""
    _FILE_CACHE.pop("bytes", None)
    _FILE_CACHE.pop("name", None)
    try:
        st.session_state.pop("_cached_file_bytes", None)
        st.session_state.pop("_cached_file_name", None)
    except Exception:
        pass


def _read_file_bytes(file) -> Tuple[bytes, str]:
    """读取上传文件字节。
    优先从模块级缓存读（跨线程可见），缓存未命中时加锁从 UploadedFile 读取。
    """
    # 第一层：模块级缓存（跨线程可见）
    cached_bytes, cached_name = _get_cached_file_data()
    if cached_bytes is not None and cached_name is not None:
        return cached_bytes, cached_name

    # 第二层：缓存未命中，加锁从原始文件读取
    if file is None:
        file = st.session_state.get("uploaded_file")
    if file is None:
        raise ValueError("请先上传 Excel 或 CSV 文件。")

    name = file.name if hasattr(file, "name") else str(file)

    # 检查文件是否已被关闭
    if hasattr(file, "closed") and file.closed:
        raise ValueError(
            f"文件对象已关闭（Streamlit 可能在 rerun 时关闭了它）。请重新上传文件。"
        )

    if hasattr(file, "read"):
        with _FILE_READ_LOCK:
            try:
                file.seek(0)
                data = file.read()
                file.seek(0)
            except ValueError as e:
                raise ValueError(f"读取文件失败（文件可能已关闭或损坏）: {e}")
            except Exception as e:
                raise ValueError(f"读取文件异常: {type(e).__name__}: {e}")
        if data:
            _set_cached_file_data(data, name)
        else:
            raise ValueError(f"文件读取为空（0 字节），请确认文件内容完整。文件名：{name}")
    elif isinstance(file, (str, Path)):
        with open(file, "rb") as f:
            data = f.read()
        _set_cached_file_data(data, name)
    else:
        raise ValueError(f"无法读取文件对象（类型：{type(file).__name__}）")

    return data, name


def _excel_engine(name: str) -> str:
    return "xlrd" if str(name).lower().endswith(".xls") else "openpyxl"


def _auto_header_row(data: bytes, name: str, sheet_name=0, max_scan: int = 8) -> int:
    """自动检测真实表头所在行（跳过标题行/合并行），返回行索引。"""
    try:
        raw = pd.read_excel(
            io.BytesIO(data), sheet_name=sheet_name, header=None,
            nrows=max_scan, engine=_excel_engine(name),
        )
    except Exception:
        return 0

    scores = []
    n_cols = max(1, raw.shape[1])
    for idx in range(len(raw)):
        row = raw.iloc[idx]
        non_null = [v for v in row if not (pd.isna(v) if not isinstance(v, str) else False)]
        if not non_null:
            scores.append(0.0)
            continue
        str_count = sum(
            1 for v in non_null
            if (lambda s: bool(s) and s.lower() != "nan" and not _is_numeric_str(s))(str(v).strip())
        )
        fill_ratio = len(non_null) / n_cols
        scores.append(str_count * fill_ratio)

    return int(pd.Series(scores).idxmax()) if scores else 0


def _is_numeric_str(s: str) -> bool:
    try:
        float(s)
        return True
    except (ValueError, TypeError):
        return False


def _get_sheet_names(file=None) -> List[str]:
    data, name = _read_file_bytes(file)
    workbook = pd.ExcelFile(io.BytesIO(data), engine=_excel_engine(name))
    return workbook.sheet_names


def _read_sheet(file, sheet_name=0, header=0) -> pd.DataFrame:
    data, name = _read_file_bytes(file)
    resolved_sheet = 0 if sheet_name == "0" else sheet_name
    return pd.read_excel(io.BytesIO(data), sheet_name=resolved_sheet, header=header, engine=_excel_engine(name))


def _load_df(file, sheet_name="0", nrows=None, auto_header: bool = True) -> pd.DataFrame:
    """内部辅助：从 session_state 或路径加载 DataFrame，自动检测表头行。"""
    data, name = _read_file_bytes(file)
    lower_name = name.lower()
    sn = 0 if sheet_name == "0" else sheet_name

    if lower_name.endswith(".csv"):
        for encoding in ("utf-8-sig", "gb18030", "gbk"):
            try:
                return pd.read_csv(io.BytesIO(data), nrows=nrows, encoding=encoding)
            except UnicodeDecodeError:
                continue
        return pd.read_csv(io.BytesIO(data), nrows=nrows)

    header_row = _auto_header_row(data, name, sheet_name=sn) if auto_header else 0
    try:
        df = pd.read_excel(
            io.BytesIO(data), sheet_name=sn, nrows=nrows,
            header=header_row, engine=_excel_engine(name),
        )
    except Exception as excel_error:
        try:
            tables = pd.read_html(io.BytesIO(data), encoding="utf-8")
        except Exception:
            try:
                tables = pd.read_html(io.BytesIO(data), encoding="gb18030")
            except Exception:
                raise excel_error
        if not tables:
            raise excel_error
        df = tables[0].head(nrows) if nrows else tables[0]
    # 统一将列名转为字符串，避免数字列名导致工具调用失败
    df.columns = [str(c) for c in df.columns]
    # 去除剩余重复列名（pandas 有时对 2026xx 类列不自动加后缀）
    seen: dict = {}
    deduped = []
    for col in df.columns:
        if col not in seen:
            seen[col] = 0
            deduped.append(col)
        else:
            seen[col] += 1
            deduped.append(f"{col}_R{seen[col]}")
    df.columns = deduped
    return df


def _looks_temporal(series: pd.Series) -> bool:
    name = str(series.name or "").lower()
    if any(token in name for token in ("date", "time", "month", "day", "日期", "时间", "月份", "月", "年")):
        return True
    sample = series.dropna().astype(str).head(10)
    if sample.empty:
        return False
    date_like = sample.str.match(r"^\d{4}[-/]?\d{1,2}([-/]?\d{1,2})?$")
    month_like = sample.str.match(r"^\d{6}$")
    return bool((date_like | month_like).all())


def _pick_dimension_column(df: pd.DataFrame, numeric_cols: List[str]) -> Optional[str]:
    candidates = [col for col in df.columns if col not in numeric_cols]
    for col in candidates:
        nunique = df[col].nunique(dropna=True)
        if 1 < nunique <= 50:
            return col
    return candidates[0] if candidates else None


def _prepare_plot_data(
    df: pd.DataFrame,
    x_column: Optional[str],
    y_column: str,
    chart_type: str,
    max_points: int = 20,
) -> Tuple[pd.Series, pd.Series, str]:
    plot_df = df.copy()
    plot_df[y_column] = pd.to_numeric(plot_df[y_column], errors="coerce")
    plot_df = plot_df.dropna(subset=[y_column])
    if plot_df.empty:
        raise ValueError(f"列 '{y_column}' 没有可绘图的数值数据。")

    if x_column and x_column in plot_df.columns:
        plot_df[x_column] = plot_df[x_column].astype(str)
        plot_df = plot_df[[x_column, y_column]]
        if plot_df[x_column].duplicated().any():
            plot_df = plot_df.groupby(x_column, as_index=False)[y_column].sum()
        if chart_type == "line":
            plot_df = plot_df.sort_values(x_column).head(max_points)
        elif len(plot_df) > max_points:
            plot_df = plot_df.nlargest(max_points, y_column)
        x_data = plot_df[x_column].reset_index(drop=True)
        x_label = x_column
    else:
        plot_df = plot_df.head(max_points).reset_index(drop=True)
        x_data = pd.Series(range(1, len(plot_df) + 1), name="行号")
        x_label = "行号"

    y_data = plot_df[y_column].reset_index(drop=True)
    return x_data, y_data, x_label


def _create_chart(
    df: pd.DataFrame,
    y_column: str,
    x_column: Optional[str] = None,
    chart_type: str = "auto",
    title: Optional[str] = None,
):
    resolved_type = chart_type
    if resolved_type == "auto":
        if x_column and x_column in df.columns and _looks_temporal(df[x_column]):
            resolved_type = "line"
        else:
            resolved_type = "bar"

    fig, ax = plt.subplots(figsize=(10, 4.8))
    if resolved_type in ("bar", "line"):
        x_data, y_data, x_label = _prepare_plot_data(df, x_column, y_column, resolved_type)
        if resolved_type == "bar":
            ax.bar(x_data.astype(str), y_data)
        else:
            ax.plot(x_data.astype(str), y_data, marker="o")
        ax.set_xlabel(x_label)
        ax.set_ylabel(y_column)
        plt.setp(ax.get_xticklabels(), rotation=45, ha="right")
    elif resolved_type == "pie":
        x_data, y_data, _ = _prepare_plot_data(df, x_column, y_column, "bar", max_points=8)
        ax.pie(y_data, labels=x_data.astype(str), autopct="%1.1f%%")
    elif resolved_type == "scatter":
        if not x_column or x_column not in df.columns:
            raise ValueError("散点图需要提供有效的 X 轴列。")
        scatter_df = df[[x_column, y_column]].copy()
        scatter_df[x_column] = pd.to_numeric(scatter_df[x_column], errors="coerce")
        scatter_df[y_column] = pd.to_numeric(scatter_df[y_column], errors="coerce")
        scatter_df = scatter_df.dropna().head(1000)
        if scatter_df.empty:
            raise ValueError("散点图需要 X/Y 均为数值列。")
        x_data = scatter_df[x_column].reset_index(drop=True)
        y_data = scatter_df[y_column].reset_index(drop=True)
        ax.scatter(x_data, y_data)
        ax.set_xlabel(x_column)
        ax.set_ylabel(y_column)
    else:
        raise ValueError(f"不支持的图表类型：{resolved_type}，可用：bar / line / pie / scatter")

    ax.set_title(title or f"{y_column}图表")
    plt.tight_layout()
    return fig, resolved_type, x_data, y_data


def _fig_to_png_bytes(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    buf.seek(0)
    return buf.getvalue()


def _summarize_chart(x_data: pd.Series, y_data: pd.Series, y_column: str, chart_type: str) -> str:
    if y_data.empty:
        return f"{y_column}暂无可解读的数据点。"
    max_idx = int(y_data.idxmax())
    min_idx = int(y_data.idxmin())
    max_x = x_data.iloc[max_idx] if max_idx < len(x_data) else "未知"
    min_x = x_data.iloc[min_idx] if min_idx < len(x_data) else "未知"
    max_y = float(y_data.iloc[max_idx])
    min_y = float(y_data.iloc[min_idx])
    mean_y = float(y_data.mean())
    if chart_type == "line":
        return f"{y_column}以趋势图展示，峰值出现在 {max_x}（{max_y:.2f}），低点出现在 {min_x}（{min_y:.2f}），平均值约 {mean_y:.2f}。"
    if chart_type == "pie":
        return f"{y_column}以占比图展示，最大扇区为 {max_x}（{max_y:.2f}），最小扇区为 {min_x}（{min_y:.2f}）。"
    if chart_type == "scatter":
        return f"{y_column}以散点图展示，共 {len(y_data)} 个点，数值范围约 {min_y:.2f} 到 {max_y:.2f}。"
    return f"{y_column}以柱状图展示，最高项为 {max_x}（{max_y:.2f}），最低项为 {min_x}（{min_y:.2f}），平均值约 {mean_y:.2f}。"


def _store_generated_chart(fig, title: str, description: str):
    st.session_state["last_plot"] = fig
    st.session_state["last_plot_title"] = title
    entry = {
        "title": title,
        "description": description,
        "image": _fig_to_png_bytes(fig),
    }
    charts = st.session_state.get("generated_charts", [])
    charts = [chart for chart in charts if chart.get("title") != title]
    charts.append(entry)
    st.session_state["generated_charts"] = charts
    return entry


_CHART_REF_PAT = re.compile(
    r"(如图\s*\d+|图\s*\d+|如下图所示|下图所示|见下图|见图\s*\d+|chart\s*\d+)",
    re.IGNORECASE,
)
_FINDING_HEADING_PAT = re.compile(r"^#{2,3}\s*(发现|核心发现|关键发现|结论)", re.IGNORECASE)


def _markdown_sections(md_text: str) -> List[str]:
    sections = [section.strip() for section in re.split(r"\n\s*\n", md_text or "") if section.strip()]
    return sections


def _section_expects_chart(section: str) -> bool:
    if _CHART_REF_PAT.search(section):
        return True
    first_line = section.splitlines()[0].strip() if section.splitlines() else ""
    return bool(_FINDING_HEADING_PAT.match(first_line))


def _interleave_report_blocks(report_markdown: str, charts: List[Dict[str, object]]) -> List[Dict[str, object]]:
    blocks: List[Dict[str, object]] = []
    chart_idx = 0
    for section in _markdown_sections(report_markdown):
        blocks.append({"type": "markdown", "content": section})
        if chart_idx < len(charts) and _section_expects_chart(section):
            chart = dict(charts[chart_idx])
            chart["label"] = f"图{chart_idx + 1}"
            blocks.append({"type": "chart", "chart": chart})
            chart_idx += 1

    if chart_idx < len(charts):
        blocks.append({"type": "markdown", "content": "## 图表补充"})
        while chart_idx < len(charts):
            chart = dict(charts[chart_idx])
            chart["label"] = f"图{chart_idx + 1}"
            blocks.append({"type": "chart", "chart": chart})
            chart_idx += 1
    return blocks


def _strip_markdown_tables(md_text: str) -> str:
    """移除 Markdown 文本中的所有表格块（以 | 开头的连续行）。"""
    lines = md_text.split("\n")
    out = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("|"):
            # 跳过整个表格块
            while i < len(lines) and lines[i].strip().startswith("|"):
                i += 1
            # 也跳过紧跟的空行
            while i < len(lines) and lines[i].strip() == "":
                i += 1
            continue
        out.append(line)
        i += 1
    return "\n".join(out)


def _filter_report_markdown(md_text: str) -> str:
    """根据侧边栏偏好过滤 Markdown 文本（去表格/去图表引用）。"""
    if not st.session_state.get("include_tables", True):
        md_text = _strip_markdown_tables(md_text)
    # 图表引用过滤：去掉形如 "（见图X）" "[图X]" "(图X)" 等引用标记
    if not st.session_state.get("include_figures", True):
        md_text = re.sub(r'[（(]?[见参]?图\s*[一二三四五六七八九十\d]+[）)]?', '', md_text)
        md_text = re.sub(r'\[图\s*[一二三四五六七八九十\d]+\]', '', md_text)
    return md_text


def _normalize_id(value) -> Optional[str]:
    if pd.isna(value):
        return None
    text = str(value).strip()
    if not text or text.lower() == "nan" or text in {"总计", "合计"}:
        return None
    try:
        numeric = float(text)
        if numeric.is_integer():
            return str(int(numeric))
    except Exception:
        pass
    return text


def _normalize_month(value) -> Optional[int]:
    if pd.isna(value):
        return None
    try:
        number = int(float(value))
    except Exception:
        match = re.search(r"(20\d{4})", str(value))
        number = int(match.group(1)) if match else None
    if number and 200001 <= number <= 209912:
        return number
    return None


def _month_label(month: int) -> str:
    text = str(int(month))
    return f"{text[2:4]}-{text[4:6]}"


def _fmt_wan(value: float) -> str:
    return f"{value:,.2f}"


def _fmt_delta(value: float) -> str:
    return f"{value:+,.2f}"


def _safe_ratio(part: float, whole: float) -> float:
    if not whole:
        return 0.0
    return abs(part) / abs(whole)


def _find_header_row(raw_df: pd.DataFrame, labels: Tuple[str, ...]) -> Tuple[int, int]:
    for row_idx in range(min(8, len(raw_df))):
        row = raw_df.iloc[row_idx]
        for col_idx, value in row.items():
            if str(value).strip() in labels:
                return row_idx, col_idx
    raise ValueError(f"未找到表头标识：{', '.join(labels)}")


def _find_column(columns, candidates: List[str], required: bool = True):
    stripped = {str(col).strip(): col for col in columns}
    for candidate in candidates:
        if candidate in stripped:
            return stripped[candidate]
    for col in columns:
        text = str(col).strip()
        if any(candidate in text for candidate in candidates):
            return col
    if required:
        raise ValueError(f"未找到列：{candidates}")
    return None


def _parse_month_matrix(file, sheet_name: str, id_labels: Tuple[str, ...], value_name: str) -> pd.DataFrame:
    raw = _read_sheet(file, sheet_name=sheet_name, header=None)
    header_row, id_col = _find_header_row(raw, id_labels)

    month_columns = []
    for col_idx, value in raw.iloc[header_row].items():
        month = _normalize_month(value)
        if month is not None:
            month_columns.append((col_idx, month))
    if not month_columns:
        raise ValueError(f"{sheet_name} 未识别到月份列。")

    records = []
    for _, row in raw.iloc[header_row + 1 :].iterrows():
        serv_id = _normalize_id(row.iloc[id_col])
        if not serv_id:
            continue
        for col_idx, month in month_columns:
            numeric = pd.to_numeric(pd.Series([row.iloc[col_idx]]), errors="coerce").iloc[0]
            if pd.isna(numeric):
                continue
            records.append({"serv_id": serv_id, "month": month, value_name: float(numeric)})

    if not records:
        raise ValueError(f"{sheet_name} 未解析到有效数据。")
    return pd.DataFrame(records)


def _parse_province_sheet(file) -> pd.DataFrame:
    province_df = _read_sheet(file, sheet_name="所属省分", header=0)
    province_col = _find_column(province_df.columns, ["所属省份", "省份"])
    id_col = _find_column(province_df.columns, ["计费ID", "SERV_ID"])
    result = province_df[[province_col, id_col]].copy()
    result.columns = ["province", "serv_id"]
    result["serv_id"] = result["serv_id"].map(_normalize_id)
    result["province"] = result["province"].astype(str).str.strip()
    result = result.dropna(subset=["serv_id"])
    result = result[result["serv_id"].astype(str) != "None"]
    return result.drop_duplicates("serv_id")


def _parse_detail_sheet(file) -> Tuple[pd.DataFrame, pd.DataFrame]:
    detail = _read_sheet(file, sheet_name="出账表明细", header=0)
    id_col = _find_column(detail.columns, ["SERV_ID计费", "SERV_ID"])
    month_col = _find_column(detail.columns, ["报收月2", "报收月"])
    amount_col = _find_column(detail.columns, ["人民币金额(按本位币)-万元", "人民币金额"])
    charge_col = _find_column(detail.columns, ["资费项名称", "账务产品名称", "产品名称"], required=False)

    parsed = pd.DataFrame(
        {
            "serv_id": detail[id_col].map(_normalize_id),
            "month": detail[month_col].map(_normalize_month),
            "rev_bill": pd.to_numeric(detail[amount_col], errors="coerce"),
            "charge_item": detail[charge_col].astype(str).str.strip() if charge_col else "未分类",
        }
    )
    parsed = parsed.dropna(subset=["serv_id", "month", "rev_bill"])
    parsed["month"] = parsed["month"].astype(int)
    aggregated = parsed.groupby(["serv_id", "month"], as_index=False)["rev_bill"].sum()
    return parsed, aggregated


def _is_reconciliation_workbook(file=None) -> bool:
    try:
        _, name = _read_file_bytes(file)
        if not str(name).lower().endswith((".xls", ".xlsx")):
            return False
        sheet_names = set(_get_sheet_names(file))
    except Exception:
        return False
    expected = {"结算前", "结算后", "所属省分", "出账表明细"}
    return expected.issubset(sheet_names)


def _classify_missing_trend(base_value: float, compare_value: float) -> str:
    if base_value <= 0 and compare_value > 0:
        return "2026 年新增/恢复出账"
    if base_value > 0 and compare_value <= 0:
        return "2026 年停止出账"
    delta = compare_value - base_value
    if abs(delta) < 1:
        return "基本持平"
    if delta > 0:
        return "收入上升"
    return "收入下降"


def _format_report_table(df: pd.DataFrame, month_cols: Optional[List[str]] = None) -> pd.DataFrame:
    table = df.copy()
    if month_cols:
        for col in month_cols:
            if col in table.columns:
                table[col] = table[col].astype(int).map(_month_label)
    float_cols = table.select_dtypes(include="number").columns
    table[float_cols] = table[float_cols].round(2)
    return table


def _build_reconciliation_analysis(file) -> Dict[str, object]:
    settle = _parse_month_matrix(file, "结算前", ("计费ID",), "settle")
    post = _parse_month_matrix(file, "结算后", ("SERV_ID计费",), "rev_post")
    province = _parse_province_sheet(file)
    detail_raw, bill = _parse_detail_sheet(file)

    settle_ids = set(settle["serv_id"])
    post_ids = set(post["serv_id"])

    monthly = (
        settle.groupby("month", as_index=False)["settle"]
        .sum()
        .merge(
            bill[bill["serv_id"].isin(settle_ids)]
            .groupby("month", as_index=False)["rev_bill"]
            .sum(),
            on="month",
            how="left",
        )
        .merge(post.groupby("month", as_index=False)["rev_post"].sum(), on="month", how="left")
    )
    monthly = monthly.fillna({"rev_bill": 0.0, "rev_post": 0.0}).sort_values("month").reset_index(drop=True)
    monthly["gap"] = monthly["rev_bill"] - monthly["settle"]
    monthly["month_label"] = monthly["month"].map(_month_label)
    monthly["year"] = monthly["month"] // 100

    years = sorted(monthly["year"].unique())
    if len(years) < 2:
        raise ValueError("对账分析至少需要两个年度。")
    base_year, compare_year = years[0], years[-1]

    yearly = monthly.groupby("year", as_index=False)[["settle", "rev_bill", "rev_post", "gap"]].sum()

    id_month = (
        settle.merge(bill, on=["serv_id", "month"], how="left")
        .fillna({"rev_bill": 0.0})
        .merge(province, on="serv_id", how="left")
    )
    id_month["province"] = id_month["province"].fillna("未匹配省分")
    id_month["gap"] = id_month["rev_bill"] - id_month["settle"]
    id_month["year"] = id_month["month"] // 100

    id_year = (
        id_month.groupby(["serv_id", "province", "year"], as_index=False)[["settle", "rev_bill", "gap"]]
        .sum()
    )

    gap_pivot = id_year.pivot_table(index=["serv_id", "province"], columns="year", values="gap", fill_value=0).reset_index()
    gap_pivot = gap_pivot.rename(columns={year: f"gap_{year}" for year in years if year in gap_pivot.columns})
    for year in years:
        gap_col = f"gap_{year}"
        if gap_col not in gap_pivot.columns:
            gap_pivot[gap_col] = 0.0
    gap_pivot["worsen"] = gap_pivot[f"gap_{compare_year}"] - gap_pivot[f"gap_{base_year}"]
    worsen_top = gap_pivot.sort_values("worsen").head(10).reset_index(drop=True)

    missing_ids = sorted(settle_ids - post_ids)
    if missing_ids:
        missing_table = (
            id_year[id_year["serv_id"].isin(missing_ids)]
            .pivot_table(index=["serv_id", "province"], columns="year", values="rev_bill", fill_value=0)
            .reset_index()
        )
        missing_table = missing_table.rename(columns={year: f"rev_{year}" for year in years if year in missing_table.columns})
        for year in years:
            col = f"rev_{year}"
            if col not in missing_table.columns:
                missing_table[col] = 0.0
        missing_table["trend"] = missing_table.apply(
            lambda row: _classify_missing_trend(row[f"rev_{base_year}"], row[f"rev_{compare_year}"]),
            axis=1,
        )
        missing_table = missing_table.sort_values(
            [f"rev_{compare_year}", f"rev_{base_year}"], ascending=False
        ).reset_index(drop=True)
    else:
        missing_table = pd.DataFrame(columns=["serv_id", "province", f"rev_{base_year}", f"rev_{compare_year}", "trend"])

    key_row = worsen_top.iloc[0] if not worsen_top.empty else None
    key_serv_id = str(key_row["serv_id"]) if key_row is not None else ""
    key_circuit = id_month[id_month["serv_id"] == key_serv_id].sort_values("month").reset_index(drop=True)

    negative_items = pd.DataFrame(columns=["charge_item", "rev_bill"])
    worst_bill_month = None
    if not key_circuit.empty:
        worst_bill_month = int(key_circuit.loc[key_circuit["rev_bill"].idxmin(), "month"])
        negative_items = (
            detail_raw[
                (detail_raw["serv_id"] == key_serv_id)
                & (detail_raw["month"] == worst_bill_month)
                & (detail_raw["rev_bill"] < 0)
            ]
            .groupby("charge_item", as_index=False)["rev_bill"]
            .sum()
            .sort_values("rev_bill")
            .reset_index(drop=True)
        )

    settle_peak = key_circuit.loc[key_circuit["settle"].idxmax()] if not key_circuit.empty else None
    settle_baseline = float(key_circuit["settle"].median()) if not key_circuit.empty else 0.0

    yearly_map = yearly.set_index("year")
    base_gap = float(yearly_map.loc[base_year, "gap"])
    compare_gap = float(yearly_map.loc[compare_year, "gap"])
    total_worsen = compare_gap - base_gap
    revenue_change = float(yearly_map.loc[compare_year, "rev_bill"] - yearly_map.loc[base_year, "rev_bill"])
    settle_change = float(yearly_map.loc[compare_year, "settle"] - yearly_map.loc[base_year, "settle"])

    worst_month = monthly.loc[monthly["gap"].idxmin()]
    top_contributor = float(worsen_top.iloc[0]["worsen"]) if not worsen_top.empty else 0.0
    top_ratio = _safe_ratio(top_contributor, total_worsen)

    active_missing_count = int((missing_table.get(f"rev_{compare_year}", pd.Series(dtype=float)) > 0).sum())
    missing_compare_total = float(missing_table.get(f"rev_{compare_year}", pd.Series(dtype=float)).sum()) if not missing_table.empty else 0.0

    yearly_table = _format_report_table(yearly.rename(columns={"year": "年份"}))
    monthly_table = _format_report_table(monthly[["month", "settle", "rev_bill", "rev_post", "gap"]], month_cols=["month"]).rename(columns={"month": "月份"})
    worsen_table = _format_report_table(
        worsen_top[["serv_id", "province", f"gap_{base_year}", f"gap_{compare_year}", "worsen"]]
    ).rename(
        columns={
            "serv_id": "计费ID",
            "province": "省份",
            f"gap_{base_year}": f"{base_year}年差额",
            f"gap_{compare_year}": f"{compare_year}年差额",
            "worsen": "恶化量",
        }
    )
    key_table = _format_report_table(key_circuit[["month", "settle", "rev_bill", "gap"]], month_cols=["month"]).rename(
        columns={"month": "月份", "settle": "省分结算", "rev_bill": "客户出账", "gap": "差额"}
    )
    missing_display = _format_report_table(missing_table).rename(
        columns={
            "serv_id": "计费ID",
            "province": "省份",
            f"rev_{base_year}": f"{base_year}年出账",
            f"rev_{compare_year}": f"{compare_year}年出账",
            "trend": "趋势判断",
        }
    )
    negative_display = _format_report_table(negative_items).rename(columns={"charge_item": "资费项名称", "rev_bill": "金额"})

    findings = [
        (
            f"{compare_year} 年客户实际出账较 {base_year} 年减少 {_fmt_wan(abs(revenue_change))} 万元，"
            f"但省分结算反而增加 {_fmt_wan(settle_change)} 万元，导致差额从 {_fmt_wan(base_gap)} 万元扩大到 {_fmt_wan(compare_gap)} 万元。"
        ),
        (
            f"利润恶化最明显的电路是 {key_serv_id}（{key_row['province']}），恶化 {_fmt_wan(abs(top_contributor))} 万元，"
            f"约占整体恶化的 {top_ratio * 100:.1f}%。"
        ) if key_row is not None else "未识别到利润恶化最明显的电路。",
        (
            f"“结算后” 口径少了 {len(missing_ids)} 条电路，其中 {active_missing_count} 条在 {compare_year} 年仍有实际出账，"
            f"合计 {_fmt_wan(missing_compare_total)} 万元。"
        ),
    ]

    anomaly_notes = []
    if worst_bill_month is not None and not negative_items.empty:
        anomaly_notes.append(
            f"{key_serv_id} 在 {_month_label(worst_bill_month)} 存在负向出账冲销，主要负项包括："
            + "；".join(
                f"{row['charge_item']} {_fmt_wan(abs(float(row['rev_bill'])))} 万元"
                for _, row in negative_items.head(3).iterrows()
            )
            + "。"
        )
    if settle_peak is not None and settle_baseline:
        spike_ratio = float(settle_peak["settle"]) / settle_baseline if settle_baseline else 0.0
        if spike_ratio >= 2:
            anomaly_notes.append(
                f"{key_serv_id} 在 {_month_label(int(settle_peak['month']))} 的省分结算升至 {_fmt_wan(float(settle_peak['settle']))} 万元，"
                f"约为常态水平的 {spike_ratio:.1f} 倍。"
            )

    recommendations = [
        f"优先核查电路 {key_serv_id} 在 {_month_label(worst_bill_month)} 的冲销依据和业务说明。" if worst_bill_month else "优先核查利润恶化最明显电路的异常月份。",
        "核对 “结算后” 口径缺失电路的映射规则，确认是否存在漏挂或收入被记为 0 的情况。",
        "把对账流程固定为：结算前成本、客户实际出账、结算后映射收入三套口径并行复核，避免仅凭单一口径下结论。",
    ]

    return {
        "base_year": base_year,
        "compare_year": compare_year,
        "monthly": monthly,
        "yearly_table": yearly_table,
        "monthly_table": monthly_table,
        "worsen_table": worsen_table,
        "key_table": key_table,
        "missing_table": missing_display,
        "negative_table": negative_display,
        "findings": findings,
        "anomaly_notes": anomaly_notes,
        "recommendations": recommendations,
        "key_serv_id": key_serv_id,
        "key_province": key_row["province"] if key_row is not None else "未匹配省分",
        "key_worsen": top_contributor,
        "top_ratio": top_ratio,
        "missing_count": len(missing_ids),
        "active_missing_count": active_missing_count,
        "missing_compare_total": missing_compare_total,
        "worst_month": worst_month,
        "total_worsen": total_worsen,
        "revenue_change": revenue_change,
        "settle_change": settle_change,
        "base_gap": base_gap,
        "compare_gap": compare_gap,
        "key_circuit": key_circuit,
    }


def _chart_overall_comparison(analysis: Dict[str, object]):
    monthly = analysis["monthly"]
    labels = monthly["month_label"].tolist()
    fig, ax = plt.subplots(figsize=(10.5, 4.8))
    ax.plot(labels, monthly["settle"], marker="o", linewidth=2, label="省分结算")
    ax.plot(labels, monthly["rev_bill"], marker="o", linewidth=2, label="客户实际出账")
    ax.plot(labels, monthly["rev_post"], marker="o", linewidth=2, label="结算后口径")
    ax.set_ylabel("万元")
    ax.set_title("月度三方口径对比")
    ax.legend()
    plt.setp(ax.get_xticklabels(), rotation=45, ha="right")
    plt.tight_layout()

    description = (
        f"{analysis['compare_year']} 年客户实际出账较 {analysis['base_year']} 年变化 {_fmt_delta(analysis['revenue_change'])} 万元，"
        f"省分结算变化 {_fmt_delta(analysis['settle_change'])} 万元，"
        f"差额由 {_fmt_wan(analysis['base_gap'])} 万元扩大到 {_fmt_wan(analysis['compare_gap'])} 万元。"
    )
    return fig, "月度三方口径对比", description


def _chart_gap_trend(analysis: Dict[str, object]):
    monthly = analysis["monthly"]
    fig, ax = plt.subplots(figsize=(10.5, 4.6))
    colors = ["#d62728" if value < 0 else "#2ca02c" for value in monthly["gap"]]
    ax.bar(monthly["month_label"], monthly["gap"], color=colors)
    ax.axhline(0, color="black", linewidth=1)
    ax.set_ylabel("万元")
    ax.set_title("月度差额（客户出账 - 省分结算）")
    plt.setp(ax.get_xticklabels(), rotation=45, ha="right")
    plt.tight_layout()

    worst_month = analysis["worst_month"]
    description = (
        f"最差月份是 {worst_month['month_label']}，当月客户出账 {_fmt_wan(float(worst_month['rev_bill']))} 万元，"
        f"较省分结算少 {_fmt_wan(abs(float(worst_month['gap'])))} 万元。"
    )
    return fig, "月度差额趋势", description


def _chart_worsen_top(analysis: Dict[str, object]):
    table = analysis["worsen_table"].copy()
    if table.empty:
        raise ValueError("暂无可用于绘图的恶化数据。")
    fig, ax = plt.subplots(figsize=(10.5, 5.2))
    labels = (table["计费ID"].astype(str) + " / " + table["省份"].astype(str)).tolist()
    ax.barh(labels[::-1], table["恶化量"].astype(float).iloc[::-1], color="#d62728")
    ax.set_xlabel("恶化量（万元，负值越大越差）")
    ax.set_title("利润恶化最大的 10 条电路")
    plt.tight_layout()

    description = (
        f"恶化最明显的是 {analysis['key_serv_id']}（{analysis['key_province']}），"
        f"恶化 {_fmt_wan(abs(analysis['key_worsen']))} 万元，约占整体恶化的 {analysis['top_ratio'] * 100:.1f}%。"
    )
    return fig, "利润恶化 Top10 电路", description


def _chart_key_circuit(analysis: Dict[str, object]):
    key_circuit = analysis["key_circuit"].copy()
    if key_circuit.empty:
        raise ValueError("关键电路无可用数据。")
    fig, ax = plt.subplots(figsize=(10.5, 4.8))
    labels = key_circuit["month"].map(_month_label).tolist()
    x = list(range(len(labels)))
    settle = pd.to_numeric(key_circuit["settle"], errors="coerce").tolist()
    rev_bill = pd.to_numeric(key_circuit["rev_bill"], errors="coerce").tolist()
    ax.plot(x, settle, marker="o", linewidth=2, label="省分结算")
    ax.plot(x, rev_bill, marker="o", linewidth=2, label="客户出账")
    ax.set_xticks(x)
    ax.set_xticklabels(labels, rotation=45, ha="right")
    ax.set_ylabel("万元")
    ax.set_title(f"关键电路 {analysis['key_serv_id']} 对比")
    ax.legend()
    plt.tight_layout()

    worst_row = key_circuit.loc[key_circuit["gap"].idxmin()]
    description = (
        f"{analysis['key_serv_id']} 在 {_month_label(int(worst_row['month']))} 的差额最差，"
        f"客户出账 {_fmt_wan(float(worst_row['rev_bill']))} 万元，对应差额 {_fmt_wan(float(worst_row['gap']))} 万元。"
    )
    return fig, f"关键电路 {analysis['key_serv_id']} 明细", description


def _chart_missing_revenue(analysis: Dict[str, object]):
    missing = analysis["missing_table"].copy()
    if missing.empty:
        raise ValueError("没有漏挂电路可绘图。")

    base_col = f"{analysis['base_year']}年出账"
    compare_col = f"{analysis['compare_year']}年出账"
    fig, ax = plt.subplots(figsize=(10.5, 5.0))
    labels = missing["计费ID"].astype(str).tolist()
    positions = range(len(labels))
    ax.bar([p - 0.18 for p in positions], missing[base_col].astype(float), width=0.36, label=str(analysis["base_year"]))
    ax.bar([p + 0.18 for p in positions], missing[compare_col].astype(float), width=0.36, label=str(analysis["compare_year"]))
    ax.set_xticks(list(positions))
    ax.set_xticklabels(labels, rotation=45, ha="right")
    ax.set_ylabel("万元")
    ax.set_title("结算后漏挂电路的实际出账")
    ax.legend()
    plt.tight_layout()

    description = (
        f"“结算后” 口径少了 {analysis['missing_count']} 条电路，"
        f"其中 {analysis['active_missing_count']} 条在 {analysis['compare_year']} 年仍有实际出账，"
        f"合计 {_fmt_wan(analysis['missing_compare_total'])} 万元。"
    )
    return fig, "漏挂电路实际出账", description


def _add_doc_table(doc: Document, title: str, df: pd.DataFrame):
    doc.add_paragraph(title)
    if df.empty:
        doc.add_paragraph("暂无数据。")
        return
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(df.columns):
        table.rows[0].cells[idx].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = "" if pd.isna(value) else str(value)


def _generate_generic_report(sheet_name="0") -> str:
    df = _load_df(None, sheet_name=sheet_name)
    numeric_cols = list(df.select_dtypes(include="number").columns)
    dimension_col = _pick_dimension_column(df, numeric_cols)
    st.session_state["generated_charts"] = []

    lines = [
        "# 数据分析报告",
        "",
        "## 1. 数据概览",
        f"- 总行数：{len(df)}",
        f"- 总列数：{len(df.columns)}",
        f"- 列名：{list(df.columns)}",
        f"- 主要维度列：{dimension_col or '未识别'}",
        f"- 数值列：{numeric_cols if numeric_cols else '无'}",
        "",
        "## 2. 各列基本统计",
    ]

    num_df = df.select_dtypes(include="number")
    lines.append(num_df.describe().round(2).to_markdown() if not num_df.empty else "（无数值列）")

    lines += ["", "## 3. 空值情况"]
    null_cols = df.isna().sum()
    null_cols = null_cols[null_cols > 0]
    if null_cols.empty:
        lines.append("所有列均无空值。")
    else:
        for col, cnt in null_cols.items():
            lines.append(f"- `{col}`：{cnt} 个空值（占 {cnt / len(df) * 100:.1f}%）")

    lines += ["", "## 4. 数值异常检测（3σ原则）"]
    outlier_lines = []
    for col in num_df.columns:
        mean, std = num_df[col].mean(), num_df[col].std()
        if pd.isna(std) or std == 0:
            continue
        outliers = df[(df[col] < mean - 3 * std) | (df[col] > mean + 3 * std)]
        if not outliers.empty:
            outlier_lines.append(f"- `{col}`：发现 {len(outliers)} 个异常值（超出均值±3σ）")
    lines.extend(outlier_lines or ["未发现明显异常值。"])

    chart_titles = []
    if numeric_cols:
        for y_column in numeric_cols[:2]:
            fig, resolved_type, x_data, y_data = _create_chart(
                df,
                y_column=y_column,
                x_column=dimension_col if dimension_col in df.columns else None,
                chart_type="auto",
                title=f"{y_column}可视化",
            )
            description = _summarize_chart(x_data, y_data, y_column, resolved_type)
            chart = _store_generated_chart(fig, f"{y_column}可视化（{resolved_type}）", description)
            chart_titles.append(chart["title"])
            plt.close(fig)

    doc = Document()
    doc.add_heading("数据分析报告", 0)
    doc.add_paragraph(f"工作表：{sheet_name}")
    doc.add_paragraph(f"总行数：{len(df)}；总列数：{len(df.columns)}")
    doc.add_paragraph(f"主要维度列：{dimension_col or '未识别'}")
    doc.add_paragraph(f"数值列：{', '.join(map(str, numeric_cols)) if numeric_cols else '无'}")

    doc.add_heading("关键发现", level=1)
    for line in (outlier_lines or ["未发现明显异常值。"]):
        doc.add_paragraph(line.lstrip("- "))

    doc.add_heading("关键数据可视化", level=1)
    if st.session_state.get("generated_charts"):
        for chart in st.session_state["generated_charts"]:
            doc.add_paragraph(chart["title"])
            doc.add_picture(io.BytesIO(chart["image"]), width=Inches(6.5))
            doc.add_paragraph(chart["description"])
    else:
        doc.add_paragraph("当前数据没有可直接作图的数值列。")

    report_buffer = io.BytesIO()
    doc.save(report_buffer)
    report_buffer.seek(0)
    uploaded = st.session_state.get("uploaded_file")
    base_name = Path(uploaded.name).stem if uploaded and hasattr(uploaded, "name") else "report"
    st.session_state["generated_report_docx"] = report_buffer.getvalue()
    st.session_state["generated_report_name"] = f"{base_name}_分析报告.docx"

    lines += ["", "## 5. 已生成真实图表"]
    if chart_titles:
        lines.extend([f"- {title}" for title in chart_titles])
        lines.extend([f"  - {chart['description']}" for chart in st.session_state["generated_charts"]])
    else:
        lines.append("当前未生成图表（可能缺少合适的数值列）。")
    lines.append("")
    lines.append("已生成可下载的 Word 报告，图表为真实 PNG 图像，并附带每张图的自动描述。")
    return "\n".join(lines)


def read_excel(sheet_name="0", nrows=50, skiprows=0) -> str:
    """读取 Excel/CSV 文件的行，返回 Markdown 表格。支持跳过前N行读取文件中部或底部。"""
    try:
        df = _load_df(None, sheet_name=sheet_name)
        total = len(df)
        start = min(int(skiprows), total)
        end = min(start + int(nrows), total)
        chunk = df.iloc[start:end]
        return f"已读取第 {start+1}-{end} 行（共 {total} 行），列名：{list(df.columns)}\n\n{chunk.to_markdown(index=False)}"
    except Exception as e:
        return f"读取失败：{e}"


def summarize_excel(sheet_name="0") -> str:
    """返回文件结构概览：行列数、列类型、空值统计、数值列基本统计量。"""
    try:
        df = _load_df(None, sheet_name=sheet_name)
        lines = [f"共 {len(df)} 行 × {len(df.columns)} 列", "", "**列信息：**"]
        for col in df.columns:
            lines.append(f"- `{col}`：类型={df[col].dtype}，空值={df[col].isna().sum()}")
        lines.append("")
        lines.append("**数值列统计：**")
        num_df = df.select_dtypes(include="number")
        lines.append(num_df.describe().round(2).to_markdown() if not num_df.empty else "（无数值列）")
        return "\n".join(lines)
    except Exception as e:
        return f"概览失败：{e}"


def filter_excel(column: str, operator: str, value: str, sheet_name="0") -> str:
    """按条件筛选数据。operator 支持：> < == != contains"""
    try:
        df = _load_df(None, sheet_name=sheet_name)
        if column not in df.columns:
            return f"列 '{column}' 不存在，可用列：{list(df.columns)}"

        if operator == "contains":
            result = df[df[column].astype(str).str.contains(value, na=False)]
        elif operator == ">":
            result = df[df[column] > float(value)]
        elif operator == "<":
            result = df[df[column] < float(value)]
        elif operator == "==":
            try:
                result = df[df[column] == float(value)]
            except ValueError:
                result = df[df[column].astype(str) == value]
        elif operator == "!=":
            try:
                result = df[df[column] != float(value)]
            except ValueError:
                result = df[df[column].astype(str) != value]
        else:
            return f"不支持的运算符：{operator}，请使用 > < == != contains"

        if result.empty:
            return "筛选结果为空。"
        return f"筛选到 {len(result)} 行：\n\n{result.head(100).to_markdown(index=False)}"
    except Exception as e:
        return f"筛选失败：{e}"


def calc_excel(column: str, operation: str, sheet_name="0") -> str:
    """对指定列做统计计算。operation 支持：sum / mean / max / min / count / median / std"""
    try:
        df = _load_df(None, sheet_name=sheet_name)
        if column not in df.columns:
            return f"列 '{column}' 不存在，可用列：{list(df.columns)}"

        col = pd.to_numeric(df[column], errors="coerce")
        ops = {
            "sum": lambda c: c.sum(),
            "mean": lambda c: round(c.mean(), 4),
            "max": lambda c: c.max(),
            "min": lambda c: c.min(),
            "count": lambda c: c.count(),
            "median": lambda c: c.median(),
            "std": lambda c: round(c.std(), 4),
        }
        if operation not in ops:
            return f"不支持的操作：{operation}，可用：{list(ops.keys())}"
        return f"列 '{column}' 的 {operation} = {ops[operation](col)}"
    except Exception as e:
        return f"计算失败：{e}"


def write_excel(row_index: int, column: str, new_value: str, sheet_name="0") -> str:
    """修改指定行列的单元格值，并保存为新文件返回下载。"""
    try:
        file = st.session_state.get("uploaded_file")
        if file is None:
            return "请先上传文件。"
        df = _load_df(file, sheet_name=sheet_name)
        if column not in df.columns:
            return f"列 '{column}' 不存在，可用列：{list(df.columns)}"
        if row_index < 0 or row_index >= len(df):
            return f"行索引 {row_index} 超出范围（共 {len(df)} 行）。"

        old_val = df.at[row_index, column]
        try:
            df.at[row_index, column] = float(new_value)
        except ValueError:
            df.at[row_index, column] = new_value

        buf = io.BytesIO()
        df.to_excel(buf, index=False)
        buf.seek(0)
        st.session_state["edited_excel"] = buf
        st.session_state["edited_excel_name"] = "edited_" + (file.name if hasattr(file, "name") else "output.xlsx")
        return f"已将第 {row_index} 行、列 '{column}' 的值从 '{old_val}' 改为 '{new_value}'。文件已准备好下载。"
    except Exception as e:
        return f"写入失败：{e}"


def plot_excel(x_column: str, y_column: str, chart_type="bar", sheet_name="0") -> str:
    """生成真实图表并存入 session_state 供 Streamlit 展示。"""
    try:
        df = _load_df(None, sheet_name=sheet_name)
        if y_column not in df.columns:
            return f"列 '{y_column}' 不存在，可用列：{list(df.columns)}"
        if x_column not in df.columns:
            return f"列 '{x_column}' 不存在，可用列：{list(df.columns)}"

        if "generated_charts" not in st.session_state:
            st.session_state["generated_charts"] = []
        fig, resolved_type, x_data, y_data = _create_chart(
            df,
            y_column=y_column,
            x_column=x_column,
            chart_type=chart_type,
            title=f"{y_column} vs {x_column}",
        )
        description = _summarize_chart(x_data, y_data, y_column, resolved_type)
        _store_generated_chart(fig, f"{y_column} vs {x_column}", description)
        plt.close(fig)
        return f"已生成真实 {resolved_type} 图表（x={x_column}，y={y_column}）。图下会附带自动描述。"
    except Exception as e:
        return f"绘图失败：{e}"


def generate_report(sheet_name="0") -> str:
    """自动生成数据分析报告（Markdown + 可下载 Word），包含真实图表和图下描述。"""
    try:
        if _is_reconciliation_workbook():
            analysis = _build_reconciliation_analysis(None)
            st.session_state["generated_charts"] = []

            chart_builders = [
                _chart_overall_comparison,
                _chart_gap_trend,
                _chart_worsen_top,
                _chart_key_circuit,
                _chart_missing_revenue,
            ]
            chart_titles = []
            for builder in chart_builders:
                try:
                    fig, title, description = builder(analysis)
                except Exception:
                    continue
                chart = _store_generated_chart(fig, title, description)
                chart_titles.append(chart["title"])
                plt.close(fig)

            lines = [
                "# 多口径对账分析报告",
                "",
                "## 1. 核心结论",
            ]
            lines.extend([f"- {item}" for item in analysis["findings"]])
            lines += ["", "## 2. 整体规模对比", analysis["yearly_table"].to_markdown(index=False)]
            lines += ["", "## 3. 月度差额趋势", analysis["monthly_table"].to_markdown(index=False)]
            lines += ["", "## 4. 利润恶化来源", analysis["worsen_table"].to_markdown(index=False)]
            lines += ["", f"## 5. 关键电路 {analysis['key_serv_id']} 明细", analysis["key_table"].to_markdown(index=False)]
            if analysis["anomaly_notes"]:
                lines.extend(["", "### 异常解释"])
                lines.extend([f"- {item}" for item in analysis["anomaly_notes"]])
            if not analysis["negative_table"].empty:
                lines += ["", "### 关键负向出账项", analysis["negative_table"].to_markdown(index=False)]
            lines += ["", "## 6. 结算后漏挂电路", analysis["missing_table"].to_markdown(index=False) if not analysis["missing_table"].empty else "未发现漏挂电路。"]
            lines += ["", "## 7. 建议动作"]
            lines.extend([f"- {item}" for item in analysis["recommendations"]])
            lines += ["", "## 8. 已生成真实图表"]
            lines.extend([f"- {title}" for title in chart_titles] or ["- 当前数据未生成图表。"])
            for chart in st.session_state.get("generated_charts", []):
                lines.append(f"  - {chart['description']}")
            lines.append("")
            lines.append("已生成可下载的 Word 报告，包含真实图表、关键表格和结论建议。")

            doc = Document()
            doc.add_heading("多口径对账分析报告", 0)
            doc.add_paragraph(f"对比年度：{analysis['base_year']} vs {analysis['compare_year']}")
            doc.add_heading("核心结论", level=1)
            for item in analysis["findings"]:
                doc.add_paragraph(item)

            doc.add_heading("整体规模对比", level=1)
            _add_doc_table(doc, "年度汇总", analysis["yearly_table"])
            doc.add_heading("月度差额趋势", level=1)
            _add_doc_table(doc, "月度明细", analysis["monthly_table"])
            doc.add_heading("利润恶化来源", level=1)
            _add_doc_table(doc, "恶化 Top10 电路", analysis["worsen_table"])
            doc.add_heading(f"关键电路 {analysis['key_serv_id']} 明细", level=1)
            _add_doc_table(doc, "关键电路逐月数据", analysis["key_table"])
            if analysis["anomaly_notes"]:
                doc.add_paragraph("异常解释")
                for item in analysis["anomaly_notes"]:
                    doc.add_paragraph(item)
            if not analysis["negative_table"].empty:
                _add_doc_table(doc, "关键负向出账项", analysis["negative_table"])
            doc.add_heading("结算后漏挂电路", level=1)
            _add_doc_table(doc, "漏挂电路实际出账", analysis["missing_table"])
            doc.add_heading("建议动作", level=1)
            for item in analysis["recommendations"]:
                doc.add_paragraph(item)
            doc.add_heading("图表", level=1)
            if st.session_state.get("generated_charts"):
                for chart in st.session_state["generated_charts"]:
                    doc.add_paragraph(chart["title"])
                    doc.add_picture(io.BytesIO(chart["image"]), width=Inches(6.5))
                    doc.add_paragraph(chart["description"])
            else:
                doc.add_paragraph("当前未生成图表。")

            report_buffer = io.BytesIO()
            doc.save(report_buffer)
            report_buffer.seek(0)
            uploaded = st.session_state.get("uploaded_file")
            base_name = Path(uploaded.name).stem if uploaded and hasattr(uploaded, "name") else "report"
            st.session_state["generated_report_docx"] = report_buffer.getvalue()
            st.session_state["generated_report_name"] = f"{base_name}_多口径分析报告.docx"
            return "\n".join(lines)

        return _generate_generic_report(sheet_name=sheet_name)
    except Exception as e:
        return f"报告生成失败：{e}"


def pivot_excel(group_column: str, value_column: str, agg_op: str = "sum",
                sheet_name: str = "0", top_n: int = 20) -> str:
    """按维度列分组聚合指标列，返回 Markdown 表格。agg_op 支持 sum/mean/count/max/min。"""
    try:
        df = _load_df(None, sheet_name=sheet_name)
        if group_column not in df.columns:
            return f"列 '{group_column}' 不存在，可用列：{list(df.columns)}"
        if value_column not in df.columns:
            return f"列 '{value_column}' 不存在，可用列：{list(df.columns)}"

        ops = {"sum", "mean", "count", "max", "min"}
        if agg_op not in ops:
            return f"不支持的操作：{agg_op}，可用：{sorted(ops)}"

        df[value_column] = pd.to_numeric(df[value_column], errors="coerce")
        result = (
            df.groupby(group_column, as_index=False)[value_column]
            .agg(agg_op)
            .sort_values(value_column, ascending=False)
            .head(top_n)
        )
        result[value_column] = result[value_column].round(4)
        return (
            f"按 '{group_column}' 分组，{agg_op}('{value_column}')，前 {len(result)} 组：\n\n"
            + result.to_markdown(index=False)
        )
    except Exception as e:
        return f"分组计算失败：{e}"


def top_n_excel(sort_column: str, n: int = 10, ascending: bool = False,
                sheet_name: str = "0") -> str:
    """返回按指定列排序后的前 N 行（默认降序取最大的 N 行）。"""
    try:
        df = _load_df(None, sheet_name=sheet_name)
        if sort_column not in df.columns:
            return f"列 '{sort_column}' 不存在，可用列：{list(df.columns)}"

        df["_sort"] = pd.to_numeric(df[sort_column], errors="coerce")
        result = df.sort_values("_sort", ascending=ascending).drop(columns=["_sort"]).head(n).reset_index(drop=True)
        direction = "升序（最小）" if ascending else "降序（最大）"
        return f"按 '{sort_column}' {direction}排列的前 {len(result)} 行：\n\n{result.to_markdown(index=False)}"
    except Exception as e:
        return f"排序失败：{e}"


def get_analysis_data(sheet_name: str = "0") -> str:
    """
    返回结构化分析数据供 LLM 解读：字段列表、数据类型、样本行、数值列统计、
    可用维度列（低基数字符串列）、可用指标列（数值列）。
    LLM 读完此工具结果后，自行决定从哪些维度分析、调用哪些工具、生成什么图表。
    """
    try:
        df = _load_df(None, sheet_name=sheet_name)
        n_rows, n_cols = df.shape

        # 字段信息
        fields = []
        dim_cols = []   # 维度列（低基数字符串列，适合分组）
        metric_cols = []  # 指标列（数值列）
        for col in df.columns:
            dtype = str(df[col].dtype)
            null_count = int(df[col].isna().sum())
            null_pct = round(null_count / n_rows * 100, 1) if n_rows > 0 else 0
            unique = int(df[col].nunique(dropna=True))
            sample = df[col].dropna().iloc[:3].tolist()
            fields.append({
                "column": col, "dtype": dtype,
                "nulls": null_count, "null_pct": f"{null_pct}%",
                "unique_values": unique, "sample": sample,
            })
            if pd.api.types.is_numeric_dtype(df[col]):
                metric_cols.append(col)
            elif unique <= max(30, n_rows * 0.05):
                dim_cols.append(col)

        # 数值列统计
        numeric_stats = {}
        num_df = df[metric_cols] if metric_cols else pd.DataFrame()
        if not num_df.empty:
            desc = num_df.describe().round(2)
            for col in desc.columns:
                numeric_stats[col] = desc[col].to_dict()

        # 样本行
        sample_rows = df.head(5).to_dict(orient="records")

        result = {
            "shape": {"rows": n_rows, "columns": n_cols},
            "fields": fields,
            "dimension_columns": dim_cols,
            "metric_columns": metric_cols,
            "numeric_stats": numeric_stats,
            "sample_rows": sample_rows,
        }
        return (
            f"数据结构已加载（{n_rows} 行 × {n_cols} 列）：\n\n"
            f"```json\n{json.dumps(result, ensure_ascii=False, indent=2, default=str)}\n```\n\n"
            f"维度列（适合分组分析）：{dim_cols}\n"
            f"指标列（适合计算/排名）：{metric_cols}"
        )
    except Exception as e:
        return f"数据读取失败：{e}"


def build_report_doc(report_markdown: str) -> str:
    """
    将 LLM 撰写的 Markdown 报告文字与已生成的图表打包为 Word (.docx) 文件。
    LLM 负责写内容，此工具只负责排版打包。
    report_markdown: LLM 生成的完整报告 Markdown 文本（含标题、发现、建议等）
    """
    try:
        doc = Document()
        include_figures = st.session_state.get("include_figures", True)
        charts = st.session_state.get("generated_charts", []) if include_figures else []
        filtered_md = _filter_report_markdown(report_markdown)
        for block in _interleave_report_blocks(filtered_md, charts):
            if block["type"] == "markdown":
                for line in block["content"].splitlines():
                    stripped = line.strip()
                    if stripped.startswith("### "):
                        doc.add_heading(stripped[4:], level=3)
                    elif stripped.startswith("## "):
                        doc.add_heading(stripped[3:], level=2)
                    elif stripped.startswith("# "):
                        doc.add_heading(stripped[2:], level=1)
                    elif stripped.startswith("- ") or stripped.startswith("* "):
                        doc.add_paragraph(stripped[2:], style="List Bullet")
                    elif re.match(r"^\d+\. ", stripped):
                        doc.add_paragraph(re.sub(r"^\d+\. ", "", stripped), style="List Number")
                    elif stripped.startswith("|"):
                        doc.add_paragraph(stripped)
                    elif stripped:
                        doc.add_paragraph(stripped)
            else:
                chart = block["chart"]
                doc.add_paragraph(f"{chart['label']}：{chart['title']}")
                doc.add_picture(io.BytesIO(chart["image"]), width=Inches(6.0))
                if chart.get("description"):
                    doc.add_paragraph(chart["description"])

        buf = io.BytesIO()
        doc.save(buf)
        uploaded = st.session_state.get("uploaded_file")
        base_name = Path(uploaded.name).stem if uploaded and hasattr(uploaded, "name") else "report"
        docx_bytes = buf.getvalue()
        st.session_state["generated_report_docx"] = docx_bytes
        st.session_state["generated_report_name"] = f"{base_name}_分析报告.docx"

        # 自动保存到侧边栏配置的目录
        save_dir = Path(st.session_state.get("save_dir", str(Path.cwd())))
        save_path = save_dir / f"{base_name}_分析报告.docx"
        os.makedirs(save_dir, exist_ok=True)
        with open(save_path, "wb") as f:
            f.write(docx_bytes)
        include_tables = st.session_state.get("include_tables", True)
        notes = []
        if not include_tables:
            notes.append("不含表格")
        if not include_figures:
            notes.append("不含图表")
        extra = "（" + "、".join(notes) + "）" if notes else ""
        return f"Word 报告已生成{extra}（含 {len(charts)} 张图表），已保存至 `{save_path}`，也可在侧边栏下载。"
    except Exception as e:
        return f"Word 报告生成失败：{e}"


# ── LaTeX 专业级模板 ──────────────────────────────────────────────────────────

_LATEX_TEMPLATE = r"""\documentclass[12pt,a4paper]{{article}}

% ── 字体 ──
\usepackage{{fontspec}}
\usepackage{{xeCJK}}
\setCJKmainfont{{SimSun}}
\setCJKsansfont{{Microsoft YaHei}}
\setmainfont{{Arial}}

% ── 页面布局 ──
\usepackage{{geometry}}
\geometry{{a4paper, top=2.5cm, bottom=2.5cm, left=2.5cm, right=2.5cm}}

% ── 颜色 ──
\usepackage{{xcolor}}
\definecolor{{primary}}{{RGB}}{{0,51,102}}
\definecolor{{accent}}{{RGB}}{{0,102,153}}
\definecolor{{headerbg}}{{RGB}}{{0,51,102}}
\definecolor{{rowalt}}{{RGB}}{{240,245,250}}
\definecolor{{textmain}}{{RGB}}{{33,33,33}}
\definecolor{{textlight}}{{RGB}}{{100,100,100}}

% ─ 图形与表格 ──
\usepackage{{graphicx}}
\usepackage{{booktabs}}
\usepackage{{longtable}}
\usepackage{{array}}
\usepackage{{colortbl}}
\usepackage{{tabularx}}
\usepackage{{multirow}}
\usepackage{{float}}

% ── 排版增强 ──
\usepackage{{parskip}}
\setlength{{\parskip}}{{0.6em}}
\usepackage{{setspace}}
\setstretch{{1.35}}
\usepackage{{titlesec}}
\usepackage{{fancyhdr}}
\usepackage{{lastpage}}
\usepackage{{enumitem}}
\usepackage{{hyperref}}
\usepackage{{tocloft}}
\usepackage{{etoolbox}}
\usepackage{{amsmath}}

% ─ 标题格式 ──
\titleformat{{\section}}
  {{\Large\bfseries\color{{primary}}}}
  {{\thesection}}
  {{1em}}
  {{}}
  [{{\color{{primary}}\titlerule[1.5pt]}}]

\titleformat{{\subsection}}
  {{\large\bfseries\color{{accent}}}}
  {{\thesubsection}}
  {{1em}}
  {{}}

\titleformat{{\subsubsection}}
  {{\normalsize\bfseries\color{{textmain}}}}
  {{\thesubsubsection}}
  {{1em}}
  {{}}

% ── 页眉页脚 ──
\pagestyle{{fancy}}
\fancyhf{{}}
\fancyhead[L]{{\small\color{{textlight}}\leftmark}}
\fancyhead[R]{{\small\color{{textlight}}\thepage}}
\fancyfoot[C]{{\small\color{{textlight}}— \thepage\ / \pageref{{LastPage}} —}}
\renewcommand{{\headrulewidth}}{{0.4pt}}
\renewcommand{{\footrulewidth}}{{0pt}}

% ── 超链接 ──
\hypersetup{{
  colorlinks=true,
  linkcolor=accent,
  urlcolor=accent,
  citecolor=accent,
  pdftitle={{分析报告}},
  pdfauthor={{公司智能助手}}
}}

% ── 图表编号 ──
\usepackage{{caption}}
\captionsetup[figure]{{font=small,labelfont=bf,labelsep=colon,textfont=it}}
\captionsetup[table]{{font=small,labelfont=bf,labelsep=colon,textfont=it}}
\numberwithin{{figure}}{{section}}
\numberwithin{{table}}{{section}}

% ─ 列表 ──
\setlist[itemize]{{leftmargin=1.5em, itemsep=0.3em, parsep=0.2em}}
\setlist[enumerate]{{leftmargin=1.5em, itemsep=0.3em, parsep=0.2em}}

% ── 封面页 ──
\newcommand{{\makecoverpage}}{{
  \begin{{titlepage}}
    \centering
    \vspace*{{3cm}}
    {cover_content}
    \vfill
    \begin{{center}}
      \large\color{{textlight}} 公司智能助手 · 自动生成\\
      \large\color{{textlight}} {report_date}
    \end{{center}}
    \vspace*{{2cm}}
  \end{{titlepage}}
}}

\begin{{document}}
\makecoverpage
\tableofcontents
\newpage
{body}
\end{{document}}
"""


def _latex_from_markdown_block(md_text: str) -> List[str]:
    """将 Markdown 文本块转换为 LaTeX 代码行。"""
    lines = md_text.splitlines()
    out = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("### "):
            out.append(r"\subsubsection{" + _latex_escape(stripped[4:]) + "}")
        elif stripped.startswith("## "):
            out.append(r"\subsection{" + _latex_escape(stripped[3:]) + "}")
        elif stripped.startswith("# "):
            out.append(r"\section{" + _latex_escape(stripped[2:]) + "}")
        elif stripped.startswith("- ") or stripped.startswith("* "):
            items = []
            while i < len(lines) and (lines[i].strip().startswith("- ") or lines[i].strip().startswith("* ")):
                items.append(r"\item " + _latex_escape(lines[i].strip()[2:]))
                i += 1
            out.append(r"\begin{itemize}")
            out.extend(items)
            out.append(r"\end{itemize}")
            continue
        elif re.match(r"^\d+\. ", stripped):
            items = []
            while i < len(lines) and re.match(r"^\d+\. ", lines[i].strip()):
                items.append(r"\item " + _latex_escape(re.sub(r"^\d+\. ", "", lines[i].strip())))
                i += 1
            out.append(r"\begin{enumerate}")
            out.extend(items)
            out.append(r"\end{enumerate}")
            continue
        elif stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            out.extend(_md_table_to_latex(table_lines))
            continue
        elif stripped == "---" or stripped == "***":
            out.append(r"\vspace{0.5em}\noindent\color{primary}\rule{\linewidth}{1.5pt}\vspace{0.5em}")
        elif stripped == "":
            out.append("")
        else:
            # 普通段落：处理行内格式后输出
            out.append(_latex_escape(stripped))
        i += 1
    return out


def _md_to_latex(md_text: str, chart_assets: List[Dict[str, object]]) -> str:
    """将 Markdown 文本转换为 LaTeX body，并在相关段落后内联图表。"""
    out = []
    for block in _interleave_report_blocks(md_text, chart_assets):
        if block["type"] == "markdown":
            out.extend(_latex_from_markdown_block(block["content"]))
        else:
            chart = block["chart"]
            safe_path = str(chart["path"]).replace("\\", "/")
            caption_text = _latex_escape(f"{chart['label']}：{chart['title']}")
            out.append(r"\begin{figure}[H]")
            out.append(r"\centering")
            out.append(r"\includegraphics[width=0.92\linewidth]{" + safe_path + "}")
            out.append(r"\caption{" + caption_text + "}")
            out.append(r"\end{figure}")
            if chart.get("description"):
                out.append(r"\small\color{textlight}" + _latex_escape(str(chart["description"])))
            out.append("")
    return "\n".join(out)


def _latex_escape(text: str) -> str:
    """转义 LaTeX 特殊字符。"""
    replacements = [
        ("\\", r"\textbackslash{}"),
        ("&", r"\&"), ("%", r"\%"), ("$", r"\$"),
        ("#", r"\#"), ("_", r"\_"), ("{", r"\{"), ("}", r"\}"),
        ("~", r"\textasciitilde{}"), ("^", r"\textasciicircum{}"),
        ("<", r"\textless{}"), (">", r"\textgreater{}"),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    # Bold **text**
    text = re.sub(r"\*\*(.+?)\*\*", r"\\textbf{\1}", text)
    # Italic *text*
    text = re.sub(r"\*(.+?)\*", r"\\textit{\1}", text)
    return text


def _md_table_to_latex(table_lines: List[str]) -> List[str]:
    """Markdown 表格 → LaTeX longtable（专业级：彩色表头 + 交替行颜色）。"""
    rows = []
    for line in table_lines:
        if re.match(r"^\|\s*[-:]+", line):
            continue  # 分隔行跳过
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return []
    n_cols = max(len(r) for r in rows)
    col_spec = "l" * n_cols
    out = [
        r"\begin{longtable}{" + col_spec + "}",
        r"\toprule",
        r"\rowcolor{headerbg}",
    ]
    for ri, row in enumerate(rows):
        escaped = [_latex_escape(c) for c in row]
        line = " & ".join(escaped) + r" \\"
        if ri == 0:
            # 表头行：白色粗体文字
            line = r"\color{white}\bfseries " + line
            out.append(line)
            out.append(r"\midrule")
        else:
            # 数据行：交替背景色
            if ri % 2 == 0:
                out.append(r"\rowcolor{rowalt}")
            out.append(line)
    out += [r"\bottomrule", r"\end{longtable}"]
    return out


def build_report_pdf(report_markdown: str) -> str:
    """
    将 LLM 撰写的 Markdown 报告 + 已生成图表编译为 PDF（通过 LaTeX/xelatex）。
    专业级排版：封面页 + 目录 + 彩色表头表格 + 交替行 + 图表编号 + 页眉页脚。
    LLM 负责写内容，此工具只负责排版打包。
    report_markdown: LLM 生成的完整报告 Markdown 文本
    """
    try:
        include_figures = st.session_state.get("include_figures", True)
        charts = st.session_state.get("generated_charts", []) if include_figures else []
        filtered_md = _filter_report_markdown(report_markdown)
        tmpdir = tempfile.mkdtemp()

        # 保存图表 PNG
        chart_assets = []
        for idx, chart in enumerate(charts):
            img_path = os.path.join(tmpdir, f"chart_{idx:02d}.png")
            with open(img_path, "wb") as f:
                f.write(chart["image"])
            chart_assets.append({**chart, "path": img_path})

        # 提取报告标题作为封面标题
        report_title = "数据分析报告"
        report_subtitle = ""
        first_line = filtered_md.splitlines()[0].strip() if filtered_md.splitlines() else ""
        if first_line.startswith("# "):
            report_title = first_line[2:].strip()
        elif first_line.startswith("## "):
            report_title = first_line[3:].strip()
        # 提取副标题（第二行如果是 ## 开头）
        lines = filtered_md.splitlines()
        for li, ln in enumerate(lines):
            if li > 0 and ln.strip().startswith("## "):
                report_subtitle = ln.strip()[3:].strip()
                break

        # 生成封面内容
        cover_content = (
            r"{\Huge\bfseries\color{primary}" + _latex_escape(report_title) + r"}\par"
            r"\vspace{1em}"
            r"{\Large\color{accent}" + _latex_escape(report_subtitle) + r"}\par"
            r"\vspace{2em}"
            r"{\large\color{textlight}共 " + str(len(charts)) + r" 张图表 · 自动生成}\par"
        )

        # 生成日期
        from datetime import datetime
        report_date = datetime.now().strftime("%Y年%m月%d日")

        # 生成 LaTeX body
        body = _md_to_latex(filtered_md, chart_assets)
        tex_source = _LATEX_TEMPLATE.format(
            body=body,
            cover_content=cover_content,
            report_date=report_date,
        )

        tex_path = os.path.join(tmpdir, "report.tex")
        with open(tex_path, "w", encoding="utf-8") as f:
            f.write(tex_source)

        # 编译两遍（确保交叉引用正确）
        xelatex = r"C:\texlive\2022\bin\win32\xelatex.exe"
        for _ in range(2):
            result = subprocess.run(
                [xelatex, "-interaction=nonstopmode", "-output-directory", tmpdir, tex_path],
                capture_output=True, text=True, encoding="utf-8", timeout=120,
            )

        pdf_path = os.path.join(tmpdir, "report.pdf")
        if not os.path.exists(pdf_path):
            log = result.stdout[-2000:] + result.stderr[-1000:]
            return f"PDF 编译失败，xelatex 日志：\n{log}"

        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

        uploaded = st.session_state.get("uploaded_file")
        base_name = Path(uploaded.name).stem if uploaded and hasattr(uploaded, "name") else "report"
        st.session_state["generated_report_pdf"] = pdf_bytes
        st.session_state["generated_report_pdf_name"] = f"{base_name}_分析报告.pdf"

        # 自动保存到侧边栏配置的目录
        save_dir = Path(st.session_state.get("save_dir", str(Path.cwd())))
        save_path = save_dir / f"{base_name}_分析报告.pdf"
        os.makedirs(save_dir, exist_ok=True)
        with open(save_path, "wb") as f:
            f.write(pdf_bytes)
        include_tables = st.session_state.get("include_tables", True)
        notes = []
        if not include_tables:
            notes.append("不含表格")
        if not include_figures:
            notes.append("不含图表")
        extra = "（" + "、".join(notes) + "）" if notes else ""
        return f"PDF 报告已生成{extra}（含 {len(charts)} 张图表），已保存至 `{save_path}`，也可在侧边栏下载。"
    except subprocess.TimeoutExpired:
        return "PDF 编译超时（xelatex 超过 120 秒），请检查内容是否过长。"
    except Exception as e:
        return f"PDF 报告生成失败：{e}"


def build_report(report_markdown: str) -> str:
    """
    一键同时生成 Word (.docx) + PDF 双格式报告。
    自动读取侧边栏的 export_word / export_pdf 开关决定输出哪些格式，
    以及 include_tables / include_figures 开关控制内容。
    report_markdown: LLM 生成的完整报告 Markdown 文本
    """
    want_word = st.session_state.get("export_word", True)
    want_pdf = st.session_state.get("export_pdf", True)

    results = []
    if want_word:
        results.append(build_report_doc(report_markdown))
    if want_pdf:
        results.append(build_report_pdf(report_markdown))

    if not results:
        return "⚠️ 未选择任何导出格式，请在侧边栏勾选 Word 或 PDF。"

    return "\n".join(results)


# ================================================================
#  新增工具：异常数据检测 (detect_anomalies)
# ================================================================

def _iqr_outliers(series: pd.Series) -> dict:
    """IQR 法检测离群点。"""
    q1 = float(series.quantile(0.25))
    q3 = float(series.quantile(0.75))
    iqr = q3 - q1
    lower = q1 - 1.5 * iqr
    upper = q3 + 1.5 * iqr
    mask = (series < lower) | (series > upper)
    return {
        "method": "IQR（四分位距法）",
        "q1": round(q1, 2), "q3": round(q3, 2), "iqr": round(iqr, 2),
        "lower_bound": round(lower, 2), "upper_bound": round(upper, 2),
        "outlier_indices": series[mask].index.tolist(),
        "outlier_count": int(mask.sum()),
        "outlier_pct": round(mask.sum() / max(len(series), 1) * 100, 1),
    }


def _zscore_outliers(series: pd.Series, threshold: float = 2.0) -> dict:
    """Z-Score 法检测离群点。"""
    mean = float(series.mean())
    std = float(series.std())
    if std == 0 or pd.isna(std):
        return {"method": "Z-Score", "error": "标准差为 0，无法计算 Z-Score", "outlier_count": 0}
    z = ((series - mean) / std).abs()
    mask = z > threshold
    return {
        "method": f"Z-Score（阈值 {threshold}σ）",
        "mean": round(mean, 2), "std": round(std, 2),
        "threshold": threshold,
        "outlier_indices": z[mask].index.tolist(),
        "outlier_count": int(mask.sum()),
        "outlier_pct": round(mask.sum() / max(len(series), 1) * 100, 1),
    }


def _isolation_forest_outliers(df: pd.DataFrame, columns: list, contamination: float = 0.05) -> dict:
    """Isolation Forest 多维异常检测（需 sklearn）。"""
    try:
        from sklearn.ensemble import IsolationForest
    except ImportError:
        return {"method": "Isolation Forest", "error": "sklearn 未安装，请运行 pip install scikit-learn"}
    sub = df[columns].copy()
    for c in columns:
        sub[c] = pd.to_numeric(sub[c], errors="coerce")
    sub = sub.dropna()
    if sub.empty:
        return {"method": "Isolation Forest", "error": "所选列无有效数值数据", "outlier_count": 0}
    model = IsolationForest(contamination=contamination, random_state=42, n_estimators=100)
    preds = model.fit_predict(sub)
    outlier_mask = preds == -1
    outlier_indices = sub[outlier_mask].index.tolist()
    return {
        "method": f"Isolation Forest（contamination={contamination}）",
        "features": columns,
        "samples_used": len(sub),
        "outlier_count": int(outlier_mask.sum()),
        "outlier_pct": round(outlier_mask.sum() / len(sub) * 100, 1),
        "outlier_indices": outlier_indices,
    }


def detect_anomalies(column: str, method: str = "auto", sheet_name: str = "0",
                     threshold: float = 2.0, contamination: float = 0.05) -> str:
    """
    对指定数值列进行异常检测。支持三种方法：
      - iqr: 四分位距法（IQR × 1.5），适合有偏分布
      - zscore: Z-Score 法（默认 ±2σ），适合近似正态分布
      - isolation_forest: 多维 Isolation Forest（需 sklearn），支持多列联合异常检测
      - auto: 自动选择（先 IQR + Z-Score，有 sklearn 则追加 Isolation Forest）

    column: 要检测的数值列名
    method: 检测方法 (auto / iqr / zscore / isolation_forest)
    threshold: Z-Score 阈值，默认 2.0
    contamination: Isolation Forest 的预期异常比例，默认 0.05
    """
    try:
        df = _load_df(None, sheet_name=sheet_name)

        if column not in df.columns:
            return f"列 '{column}' 不存在，可用列：{list(df.columns)}"

        series = pd.to_numeric(df[column], errors="coerce").dropna()
        if len(series) < 4:
            return f"列 '{column}' 有效数值不足（仅 {len(series)} 个），无法做统计异常检测。"

        results = []

        if method in ("iqr", "auto"):
            iqr_result = _iqr_outliers(series)
            results.append(f"## {iqr_result['method']}")
            if "error" not in iqr_result:
                results.append(
                    f"- 四分位数：Q1={iqr_result['q1']}，Q3={iqr_result['q3']}，IQR={iqr_result['iqr']}\n"
                    f"- 异常边界：[{iqr_result['lower_bound']}, {iqr_result['upper_bound']}]\n"
                    f"- 异常点：{iqr_result['outlier_count']} 个（占 {iqr_result['outlier_pct']}%）"
                )
                if iqr_result["outlier_indices"]:
                    outlier_vals = series.iloc[iqr_result["outlier_indices"][:10]]
                    results.append(f"- 前 10 个异常值：{outlier_vals.to_dict()}")
            else:
                results.append(f"- {iqr_result['error']}")

        if method in ("zscore", "auto"):
            z_result = _zscore_outliers(series, threshold)
            results.append(f"\n## {z_result['method']}")
            if "error" not in z_result:
                results.append(
                    f"- 均值={z_result['mean']}，标准差={z_result['std']}\n"
                    f"- 超出 ±{z_result['threshold']}σ 的点：{z_result['outlier_count']} 个（占 {z_result['outlier_pct']}%）"
                )
                if z_result["outlier_indices"]:
                    outlier_vals = series.iloc[z_result["outlier_indices"][:10]]
                    results.append(f"- 前 10 个异常值：{outlier_vals.to_dict()}")
            else:
                results.append(f"- {z_result['error']}")

        if method in ("isolation_forest", "auto"):
            if_column_result = _isolation_forest_outliers(df, [column], contamination)
            results.append(f"\n## {if_column_result['method']}")
            if "error" in if_column_result:
                results.append(f"- {if_column_result['error']}")
            else:
                results.append(
                    f"- 样本数：{if_column_result['samples_used']}\n"
                    f"- 异常点：{if_column_result['outlier_count']} 个（占 {if_column_result['outlier_pct']}%）"
                )

        # 交叉汇总
        results.append("")
        results.append("---")
        results.append(f"📊 **{column} 异常检测汇总**：以上方法交叉验证，建议重点关注被多种方法同时标为异常的数据点。")

        return "\n".join(results)
    except Exception as e:
        return f"异常检测失败：{e}"


# ================================================================
#  新增工具：多口径对账 (reconcile_sheets)
# ================================================================

def reconcile_sheets(key_column: str, compare_columns: list = None,
                     sheet_a: str = "0", sheet_b: str = "1",
                     tolerance: float = 0.01) -> str:
    """
    对两个 Sheet 按关键列进行对账，找出差异。

    key_column: 两 Sheet 之间匹配用的关键列名（如 '计费ID'、'合同号'）
    compare_columns: 要比对的数值列名列表。留空则自动选两个 Sheet 的公有数值列
    sheet_a: 基准 Sheet（如 "系统数据"），默认 0
    sheet_b: 对比 Sheet（如 "台账数据"），默认 1
    tolerance: 差异容忍度（绝对值），小于此值的差异不报告，默认 0.01
    """
    try:
        df_a = _load_df(None, sheet_name=sheet_a)
        df_b = _load_df(None, sheet_name=sheet_b)

        # 检查关键列
        if key_column not in df_a.columns:
            return f"基准 Sheet '{sheet_a}' 中不存在关键列 '{key_column}'，可用列：{list(df_a.columns)}"
        if key_column not in df_b.columns:
            return f"对比 Sheet '{sheet_b}' 中不存在关键列 '{key_column}'，可用列：{list(df_b.columns)}"

        # 自动选择比对列
        if not compare_columns:
            num_a = [c for c in df_a.columns if pd.api.types.is_numeric_dtype(df_a[c]) and c != key_column]
            num_b = [c for c in df_b.columns if pd.api.types.is_numeric_dtype(df_b[c]) and c != key_column]
            compare_columns = [c for c in num_a if c in num_b]
            if not compare_columns:
                return (f"两个 Sheet 未找到公有数值列可用于比对。\n"
                        f"Sheet '{sheet_a}' 数值列：{num_a}\n"
                        f"Sheet '{sheet_b}' 数值列：{num_b}")

        # 规范化关键列
        df_a["_key"] = df_a[key_column].astype(str).str.strip()
        df_b["_key"] = df_b[key_column].astype(str).str.strip()

        # 合并
        merged = df_a[["_key"] + compare_columns].merge(
            df_b[["_key"] + compare_columns],
            on="_key", how="outer", suffixes=("_基准", "_对比"), indicator=True
        )

        # 统计
        both = int((merged["_merge"] == "both").sum())
        only_a = int((merged["_merge"] == "left_only").sum())
        only_b = int((merged["_merge"] == "right_only").sum())

        lines = [
            f"# 多口径对账结果",
            f"",
            f"**基准 Sheet**：{sheet_a}（{len(df_a)} 行）",
            f"**对比 Sheet**：{sheet_b}（{len(df_b)} 行）",
            f"**匹配关键列**：{key_column}",
            f"**比对列**：{', '.join(compare_columns)}",
            f"**差异容忍度**：±{tolerance}",
            f"",
            f"## 匹配概况",
            f"| 状态 | 数量 |",
            f"|------|------|",
            f"| 两方均存在 | {both} |",
            f"| 仅基准方有 | {only_a} |",
            f"| 仅对比方有 | {only_b} |",
        ]

        # 仅一方存在的记录
        if only_a > 0:
            missing_b = merged[merged["_merge"] == "left_only"]["_key"].head(20).tolist()
            lines.append(f"\n### 仅基准方存在（对比方缺失）的前 20 条")
            lines.append(f"```\n{', '.join(missing_b)}\n```")

        if only_b > 0:
            missing_a = merged[merged["_merge"] == "right_only"]["_key"].head(20).tolist()
            lines.append(f"\n### 仅对比方存在（基准方缺失）的前 20 条")
            lines.append(f"```\n{', '.join(missing_a)}\n```")

        # 数值差异
        lines.append(f"\n## 数值差异明细")
        diff_rows = []
        for col in compare_columns:
            col_a = f"{col}_基准"
            col_b = f"{col}_对比"
            if col_a not in merged.columns or col_b not in merged.columns:
                continue
            merged[col_a] = pd.to_numeric(merged[col_a], errors="coerce").fillna(0)
            merged[col_b] = pd.to_numeric(merged[col_b], errors="coerce").fillna(0)
            diff = (merged[col_a] - merged[col_b]).abs()
            mask = (merged["_merge"] == "both") & (diff > tolerance)
            n_diff = int(mask.sum())
            if n_diff > 0:
                sub = merged.loc[mask, ["_key", col_a, col_b]].copy()
                sub["差异"] = (sub[col_a] - sub[col_b]).round(4)
                sub = sub.rename(columns={"_key": key_column, col_a: f"{col}(基准)", col_b: f"{col}(对比)"})
                sub = sub.sort_values("差异", key=abs, ascending=False).head(30)
                lines.append(f"\n### {col}：{n_diff} 条差异记录（展示前 30）")
                lines.append(sub.to_markdown(index=False))
                diff_rows.append(n_diff)

        if not diff_rows:
            lines.append(f"\n✅ 所有 {both} 条匹配记录在容忍度 ±{tolerance} 内完全一致。")

        # 汇总
        total_diff = sum(diff_rows)
        lines.append(f"\n---")
        lines.append(f"📊 **对账汇总**：共 {both + only_a + only_b} 条唯一记录，{both} 条匹配，{total_diff} 条存在数值差异，{only_a + only_b} 条仅单方存在。")

        return "\n".join(lines)
    except Exception as e:
        return f"对账失败：{e}"
